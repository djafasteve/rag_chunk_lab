from typing import List, Dict
from pypdf import PdfReader
import re, os, json
from functools import lru_cache

# Support pour les documents Word
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import olefile
    from docx import Document as DocxDocument
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

def load_document(path: str) -> List[Dict]:
    """
    Charge un document et retourne une liste de pages avec leur texte
    Support: PDF, DOCX, DOC, TXT, MD
    """
    ext = os.path.splitext(path)[1].lower()
    pages = []

    if ext == '.pdf':
        reader = PdfReader(path)
        for i, page in enumerate(reader.pages):
            try:
                txt = page.extract_text() or ''
            except Exception:
                txt = ''
            pages.append({'page': i+1, 'text': txt, 'source_file': os.path.basename(path)})

        if not pages:
            pages = [{'page': 1, 'text': '', 'source_file': os.path.basename(path)}]

    elif ext == '.docx':
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx n'est pas installé. Installer avec: pip install python-docx")

        try:
            doc = DocxDocument(path)
            full_text = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Traiter les tableaux aussi
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            txt = '\n'.join(full_text)
            pages = [{'page': 1, 'text': txt, 'source_file': os.path.basename(path)}]

        except Exception as e:
            print(f"Erreur lecture DOCX {path}: {e}")
            pages = [{'page': 1, 'text': '', 'source_file': os.path.basename(path)}]

    elif ext == '.doc':
        # Pour les anciens fichiers .doc, utiliser une conversion basique
        print(f"⚠️  Fichier .doc détecté: {os.path.basename(path)}")
        print(f"    Conversion .doc limitée - pour de meilleurs résultats, convertir en .docx")
        try:
            # Tentative de lecture brute (peut être très imparfaite)
            with open(path, 'rb') as f:
                content = f.read()
                # Extraction basique du texte (très imparfaite)
                txt = content.decode('utf-8', errors='ignore')
                # Nettoyer un peu
                txt = re.sub(r'[\x00-\x1f\x7f-\x9f]', ' ', txt)
                txt = re.sub(r'\s+', ' ', txt)
            pages = [{'page': 1, 'text': txt, 'source_file': os.path.basename(path)}]
        except Exception as e:
            print(f"Erreur lecture DOC {path}: {e}")
            pages = [{'page': 1, 'text': '', 'source_file': os.path.basename(path)}]

    elif ext in ['.txt', '.md']:
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                txt = f.read()
            pages = [{'page': 1, 'text': txt, 'source_file': os.path.basename(path)}]
        except Exception as e:
            print(f"Erreur lecture {ext} {path}: {e}")
            pages = [{'page': 1, 'text': '', 'source_file': os.path.basename(path)}]

    else:
        print(f"⚠️  Format non supporté: {ext} pour {os.path.basename(path)}")
        pages = []

    return pages

@lru_cache(maxsize=1000)
def tokenize_words(text: str) -> List[str]:
    """OPTIMISATION: Cache LRU pour éviter la re-tokenisation du même texte"""
    return re.findall(r"\w+|\S", text, re.UNICODE)

def join_tokens(tokens: List[str]) -> str:
    out = []
    for i, t in enumerate(tokens):
        if i > 0 and not re.match(r"\W$", tokens[i-1]) and not re.match(r"\W$", t):
            out.append(' ')
        out.append(t)
    return ''.join(out)

def iter_headings(text: str):
    # Yield (start, end, title) for simple headings common in regulatory docs
    pattern = re.compile(r"(?im)^(?P<title>\s*(Article\s+\d+\b|\bSection\s+\d+\b|\bChapitre\s+\d+\b|[A-Z][A-Z0-9 \-]{4,}))\s*$")
    for m in pattern.finditer(text):
        yield m.start(), m.end(), m.group('title').strip()

def save_json(path: str, obj) -> None:
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
